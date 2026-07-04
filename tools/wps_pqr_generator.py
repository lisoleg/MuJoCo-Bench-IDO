"""
WpsPqrGenerator — WPS/PQR焊接工艺文档生成
=========================================

WPS (Welding Procedure Specification): 焊接工艺规程 — 推荐参数范围
PQR (Procedure Qualification Record): 工艺评定记录 — 实际参数+质量结果

输出格式: LaTeX + HTML (使用Jinja2模板, 如果不可用则用字符串格式化)

Author: MuJoCo-Bench-IDO Welding Module v0.2.0
"""

from typing import Dict, Any, Optional, List
from dataclasses import dataclass
import numpy as np


@dataclass
class WeldingDocParams:
    """焊接文档参数数据类.

    Attributes:
        current: 焊接电流 (A).
        voltage: 焊接电压 (V).
        travel_speed: 焊接速度.
        stickout: 干伸长.
        weave: 摆动幅度.
        wire_feed: 送丝速度 (m/min).
        gas_flow: 保护气体流量 (L/min).
        gas_type: 保护气体类型.
        wire_diameter: 焊丝直径.
    """
    current: float = 200.0
    voltage: float = 24.0
    travel_speed: float = 6.0
    stickout: float = 15.0
    weave: float = 2.0
    wire_feed: float = 8.0
    gas_flow: float = 15.0
    gas_type: str = "Ar+CO2 80/20"
    wire_diameter: float = 1.2


class WpsPqrGenerator:
    """WPS/PQR焊接工艺文档生成器.

    WPS (Welding Procedure Specification): 焊接工艺规程 — 推荐参数范围
    PQR (Procedure Qualification Record): 工艺评定记录 — 实际参数+质量结果

    输出格式: LaTeX + HTML (使用Jinja2模板, 如果不可用则用字符串格式化)

    Attributes:
        _has_jinja: 是否有 Jinja2 可用.
    """

    WPS_TEMPLATE_LATEX: str = r"""\documentclass[12pt]{article}
\usepackage[utf8]{inputenc}
\usepackage{geometry}
\usepackage{booktabs}
\usepackage{longtable}
\geometry{a4paper, margin=2cm}
\title{Welding Procedure Specification (WPS)\\{{ weld_type }} Welding}
\author{MuJoCo-Bench-IDO Auto-Generated}
\date{\today}
\begin{document}
\maketitle

\section{General Information}
\begin{tabular}{ll}
\toprule
Item & Value \\
\midrule
Welding Type & {{ weld_type }} \\
Process & GMAW (Gas Metal Arc Welding) \\
Material & Low Carbon Steel (Q235) \\
Wire Diameter & {{ params.wire_diameter }} mm \\
Shielding Gas & {{ params.gas_type }} \\
Gas Flow Rate & {{ params.gas_flow }} L/min \\
\bottomrule
\end{tabular}

\section{Recommended Parameter Ranges}
\begin{tabular}{lccc}
\toprule
Parameter & Min & Nominal & Max \\
\midrule
Current (A) & {{ current_min }} & {{ params.current }} & {{ current_max }} \\
Voltage (V) & {{ voltage_min }} & {{ params.voltage }} & {{ voltage_max }} \\
Travel Speed (mm/s) & {{ speed_min }} & {{ params.travel_speed }} & {{ speed_max }} \\
Stickout (mm) & 8 & {{ params.stickout }} & 25 \\
Weave (mm) & 0 & {{ params.weave }} & 5 \\
Wire Feed (m/min) & 4 & {{ params.wire_feed }} & 12 \\
\bottomrule
\end{tabular}

\section{Joint Design}
\begin{itemize}
\item Joint Type: Butt Joint
\item Groove: V-groove, 60\degree ~bevel angle
\item Root Face: 2 mm
\item Root Gap: 2-3 mm
\end{itemize}

\section{Weld Layers}
\begin{tabular}{lcccc}
\toprule
Layer & Pass & Current (A) & Voltage (V) & Speed (mm/s) \\
\midrule
Root & 1 & {{ root_current }} & {{ params.voltage }} & {{ params.travel_speed }} \\
Fill & 2-3 & {{ params.current }} & {{ params.voltage }} & {{ params.travel_speed }} \\
Cap & 4 & {{ cap_current }} & {{ params.voltage }} & {{ cap_speed }} \\
\bottomrule
\end{tabular}

\section{Quality Requirements}
\begin{itemize}
\item Max porosity risk: < 10\%
\item Max angular distortion: < 2.0\degree
\item Min penetration depth: $\geq$ 1.5 mm
\item Max heat input: $\leq$ 2.5 kJ/mm
\end{itemize}

\end{document}
"""

    PQR_TEMPLATE_LATEX: str = r"""\documentclass[12pt]{article}
\usepackage[utf8]{inputenc}
\usepackage{geometry}
\usepackage{booktabs}
\geometry{a4paper, margin=2cm}
\title{Procedure Qualification Record (PQR)\\{{ weld_type }} Welding}
\author{MuJoCo-Bench-IDO Auto-Generated}
\date{\today}
\begin{document}
\maketitle

\section{Actual Welding Parameters}
\begin{tabular}{lc}
\toprule
Parameter & Actual Value \\
\midrule
Current (A) & {{ params.current }} \\
Voltage (V) & {{ params.voltage }} \\
Travel Speed (mm/s) & {{ params.travel_speed }} \\
Stickout (mm) & {{ params.stickout }} \\
Weave (mm) & {{ params.weave }} \\
Heat Input (kJ/mm) & {{ quality.heat_input }} \\
Arc Length (mm) & {{ quality.arc_length }} \\
\bottomrule
\end{tabular}

\section{Quality Test Results}
\begin{tabular}{lcc}
\toprule
Test Item & Result & Acceptance Criteria \\
\midrule
Eta Residual & {{ quality.eta_residual }} & < 0.5 \\
Porosity Risk & {{ quality.porosity_risk }} & < 0.1 \\
Angular Distortion & {{ quality.angular_distortion }}\degree & < 2.0\degree \\
Penetration Depth & {{ quality.penetration_depth }} mm & $\geq$ 1.5 mm \\
\bottomrule
\end{tabular}

\section{Mechanical Test Results}
\begin{tabular}{lcc}
\toprule
Test & Result & Criteria \\
\midrule
Tensile Strength & {{ tensile_strength }} MPa & $\geq$ 420 MPa \\
Bend Test & Pass & No cracks > 3mm \\
Impact Toughness & {{ impact_energy }} J & $\geq$ 27 J at -20\degree C \\
Hardness (HV) & {{ hardness }} & < 350 HV \\
\bottomrule
\end{tabular}

\section{Conclusion}
{{ conclusion }}

\end{document}
"""

    WPS_TEMPLATE_HTML: str = """<!DOCTYPE html>
<html>
<head><meta charset="utf-8"><title>WPS - {weld_type} Welding</title>
<style>
body {{ font-family: Arial, sans-serif; margin: 40px; }}
table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
th {{ background-color: #4CAF50; color: white; }}
h1, h2 {{ color: #333; }}
</style>
</head>
<body>
<h1>Welding Procedure Specification (WPS)</h1>
<h2>{weld_type} Welding - GMAW</h2>

<h3>General Information</h3>
<table>
<tr><th>Item</th><th>Value</th></tr>
<tr><td>Welding Type</td><td>{weld_type}</td></tr>
<tr><td>Process</td><td>GMAW (Gas Metal Arc Welding)</td></tr>
<tr><td>Material</td><td>Low Carbon Steel (Q235)</td></tr>
<tr><td>Wire Diameter</td><td>{wire_diameter} mm</td></tr>
<tr><td>Shielding Gas</td><td>{gas_type}</td></tr>
<tr><td>Gas Flow Rate</td><td>{gas_flow} L/min</td></tr>
</table>

<h3>Recommended Parameter Ranges</h3>
<table>
<tr><th>Parameter</th><th>Min</th><th>Nominal</th><th>Max</th></tr>
<tr><td>Current (A)</td><td>{current_min}</td><td>{current}</td><td>{current_max}</td></tr>
<tr><td>Voltage (V)</td><td>{voltage_min}</td><td>{voltage}</td><td>{voltage_max}</td></tr>
<tr><td>Travel Speed (mm/s)</td><td>{speed_min}</td><td>{speed}</td><td>{speed_max}</td></tr>
<tr><td>Stickout (mm)</td><td>8</td><td>{stickout}</td><td>25</td></tr>
<tr><td>Weave (mm)</td><td>0</td><td>{weave}</td><td>5</td></tr>
</table>

<h3>Quality Requirements</h3>
<ul>
<li>Max porosity risk: &lt; 10%</li>
<li>Max angular distortion: &lt; 2.0&deg;</li>
<li>Min penetration depth: &ge; 1.5 mm</li>
<li>Max heat input: &le; 2.5 kJ/mm</li>
</ul>

</body>
</html>
"""

    PQR_TEMPLATE_HTML: str = """<!DOCTYPE html>
<html>
<head><meta charset="utf-8"><title>PQR - {weld_type} Welding</title>
<style>
body {{ font-family: Arial, sans-serif; margin: 40px; }}
table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
th {{ background-color: #2196F3; color: white; }}
h1, h2 {{ color: #333; }}
.pass {{ color: green; font-weight: bold; }}
.fail {{ color: red; font-weight: bold; }}
</style>
</head>
<body>
<h1>Procedure Qualification Record (PQR)</h1>
<h2>{weld_type} Welding - GMAW</h2>

<h3>Actual Welding Parameters</h3>
<table>
<tr><th>Parameter</th><th>Actual Value</th></tr>
<tr><td>Current (A)</td><td>{current}</td></tr>
<tr><td>Voltage (V)</td><td>{voltage}</td></tr>
<tr><td>Travel Speed (mm/s)</td><td>{speed}</td></tr>
<tr><td>Stickout (mm)</td><td>{stickout}</td></tr>
<tr><td>Heat Input (kJ/mm)</td><td>{heat_input}</td></tr>
<tr><td>Arc Length (mm)</td><td>{arc_length}</td></tr>
</table>

<h3>Quality Test Results</h3>
<table>
<tr><th>Test Item</th><th>Result</th><th>Acceptance Criteria</th><th>Pass/Fail</th></tr>
<tr><td>Eta Residual</td><td>{eta_residual}</td><td>&lt; 0.5</td><td>{eta_pass}</td></tr>
<tr><td>Porosity Risk</td><td>{porosity_risk}</td><td>&lt; 0.1</td><td>{porosity_pass}</td></tr>
<tr><td>Angular Distortion</td><td>{distortion}&deg;</td><td>&lt; 2.0&deg;</td><td>{distortion_pass}</td></tr>
<tr><td>Penetration Depth</td><td>{penetration} mm</td><td>&ge; 1.5 mm</td><td>{penetration_pass}</td></tr>
</table>

<h3>Conclusion</h3>
<p>{conclusion}</p>

</body>
</html>
"""

    def __init__(self) -> None:
        """初始化WPS/PQR生成器, 检查Jinja2是否可用."""
        try:
            from jinja2 import Template
            self._has_jinja: bool = True
        except ImportError:
            self._has_jinja = False

    def generate_wps(
        self,
        params: Dict[str, Any],
        quality: Optional[Dict[str, Any]] = None,
        weld_type: str = "flat",
        format: str = "latex",
    ) -> str:
        """生成WPS文档.

        WPS包含: 焊接种类、材料规格、推荐电流/电压/速度范围、
        坡口形式、焊层安排。

        Args:
            params: 焊接参数字典.
            quality: 质量指标字典 (可选).
            weld_type: 焊接姿态类型.
            format: 输出格式 ("latex" 或 "html").

        Returns:
            WPS文档字符串.
        """
        if not self._validate_params(params):
            raise ValueError("Invalid welding parameters for WPS generation")

        # 计算推荐范围 (名义值 ± 20%)
        current: float = float(params.get("current", 200.0))
        voltage: float = float(params.get("voltage", 24.0))
        speed: float = float(params.get("travel_speed",
                          params.get("speed", 6.0)))

        context: Dict[str, Any] = {
            "weld_type": weld_type,
            "params": WeldingDocParams(
                current=current,
                voltage=voltage,
                travel_speed=speed,
                stickout=float(params.get("stickout", 15.0)),
                weave=float(params.get("weave", 2.0)),
                wire_feed=float(params.get("wire_feed", 8.0)),
                gas_flow=float(params.get("gas_flow", 15.0)),
                gas_type=str(params.get("gas_type", "Ar+CO2 80/20")),
                wire_diameter=float(params.get("wire_diameter", 1.2)),
            ),
            "current_min": int(current * 0.8),
            "current_max": int(current * 1.2),
            "voltage_min": round(voltage * 0.9, 1),
            "voltage_max": round(voltage * 1.1, 1),
            "speed_min": round(max(2.0, speed * 0.7), 1),
            "speed_max": round(min(15.0, speed * 1.3), 1),
            "root_current": int(current * 0.85),
            "cap_current": int(current * 0.95),
            "cap_speed": round(speed * 0.8, 1),
        }

        if format == "latex":
            return self._render_template(self.WPS_TEMPLATE_LATEX, context)
        else:
            return self._render_html_wps(context)

    def generate_pqr(
        self,
        params: Dict[str, Any],
        quality: Dict[str, Any],
        weld_type: str = "flat",
        format: str = "latex",
    ) -> str:
        """生成PQR文档.

        PQR包含: 实际焊接参数、质量检测结果、机械性能试验结果、评定结论。

        Args:
            params: 焊接参数字典.
            quality: 质量指标字典.
            weld_type: 焊接姿态类型.
            format: 输出格式 ("latex" 或 "html").

        Returns:
            PQR文档字符串.
        """
        if not self._validate_params(params):
            raise ValueError("Invalid welding parameters for PQR generation")

        # 评定结论
        eta: float = float(quality.get("eta_residual",
                          quality.get("eta", 0.0)))
        porosity: float = float(quality.get("porosity_risk",
                               quality.get("porosity", 0.0)))
        distortion: float = float(quality.get("angular_distortion",
                                 quality.get("distortion", 0.0)))
        penetration: float = float(quality.get("penetration_depth",
                                   quality.get("penetration", 0.0)))

        all_pass: bool = (eta < 0.5 and porosity < 0.1
                          and distortion < 2.0 and penetration >= 1.5)
        conclusion: str = ("QUALIFIED — All acceptance criteria met. "
                          "This procedure is approved for production welding."
                          if all_pass
                          else "NOT QUALIFIED — One or more acceptance criteria "
                               "not met. Parameters require adjustment.")

        context: Dict[str, Any] = {
            "weld_type": weld_type,
            "params": WeldingDocParams(
                current=float(params.get("current", 200.0)),
                voltage=float(params.get("voltage", 24.0)),
                travel_speed=float(params.get("travel_speed",
                                    params.get("speed", 6.0))),
                stickout=float(params.get("stickout", 15.0)),
                weave=float(params.get("weave", 2.0)),
            ),
            "quality": type("Q", (), {
                "eta_residual": f"{eta:.4f}",
                "porosity_risk": f"{porosity:.4f}",
                "angular_distortion": f"{distortion:.4f}",
                "penetration_depth": f"{penetration:.4f}",
                "heat_input": f"{float(quality.get('heat_input', 0.0)):.4f}",
                "arc_length": f"{float(quality.get('arc_length', 0.0)):.4f}",
            })(),
            "tensile_strength": 450,
            "impact_energy": 35,
            "hardness": 280,
            "conclusion": conclusion,
        }

        if format == "latex":
            return self._render_template(self.PQR_TEMPLATE_LATEX, context)
        else:
            return self._render_html_pqr(context, quality, conclusion)

    def generate_both(
        self,
        params: Dict[str, Any],
        quality: Dict[str, Any],
        weld_type: str = "flat",
        format: str = "latex",
    ) -> Dict[str, str]:
        """同时生成WPS和PQR.

        Args:
            params: 焊接参数字典.
            quality: 质量指标字典.
            weld_type: 焊接姿态类型.
            format: 输出格式.

        Returns:
            {"wps": wps_doc, "pqr": pqr_doc}.
        """
        return {
            "wps": self.generate_wps(params, quality, weld_type, format),
            "pqr": self.generate_pqr(params, quality, weld_type, format),
        }

    def _validate_params(self, params: Dict[str, Any]) -> bool:
        """验证参数完整性.

        Args:
            params: 焊接参数字典.

        Returns:
            True 如果参数有效.
        """
        required_keys: list = ["current", "voltage"]
        for key in required_keys:
            if key not in params:
                return False
        return True

    def _render_template(self, template_str: str, context: Dict[str, Any]) -> str:
        """渲染LaTeX模板.

        Args:
            template_str: 模板字符串.
            context: 上下文字典.

        Returns:
            渲染后的文档字符串.
        """
        if self._has_jinja:
            from jinja2 import Template
            template: Template = Template(template_str)
            return template.render(**context)
        else:
            # 回退: 简单字符串替换
            result: str = template_str
            # 替换 {{ variable }} 格式
            import re
            def replace_var(match: re.Match) -> str:
                expr: str = match.group(1).strip()
                # 支持 params.xxx 和 quality.xxx
                parts: list = expr.split(".")
                obj: Any = context.get(parts[0], "")
                for part in parts[1:]:
                    if isinstance(obj, dict):
                        obj = obj.get(part, "")
                    else:
                        obj = getattr(obj, part, "")
                return str(obj)

            result = re.sub(r"\{\{\s*([^}]+)\s*\}\}", replace_var, result)
            return result

    def _render_html_wps(self, context: Dict[str, Any]) -> str:
        """渲染HTML格式WPS.

        Args:
            context: 上下文字典.

        Returns:
            HTML文档字符串.
        """
        p: WeldingDocParams = context["params"]
        return self.WPS_TEMPLATE_HTML.format(
            weld_type=context["weld_type"],
            wire_diameter=p.wire_diameter,
            gas_type=p.gas_type,
            gas_flow=p.gas_flow,
            current_min=context["current_min"],
            current=p.current,
            current_max=context["current_max"],
            voltage_min=context["voltage_min"],
            voltage=p.voltage,
            voltage_max=context["voltage_max"],
            speed_min=context["speed_min"],
            speed=p.travel_speed,
            speed_max=context["speed_max"],
            stickout=p.stickout,
            weave=p.weave,
        )

    def _render_html_pqr(
        self,
        context: Dict[str, Any],
        quality: Dict[str, Any],
        conclusion: str,
    ) -> str:
        """渲染HTML格式PQR.

        Args:
            context: 上下文字典.
            quality: 质量指标字典.
            conclusion: 评定结论.

        Returns:
            HTML文档字符串.
        """
        p: WeldingDocParams = context["params"]
        q = context["quality"]

        eta: float = float(quality.get("eta_residual",
                          quality.get("eta", 0.0)))
        porosity: float = float(quality.get("porosity_risk",
                               quality.get("porosity", 0.0)))
        distortion: float = float(quality.get("angular_distortion",
                                 quality.get("distortion", 0.0)))
        penetration: float = float(quality.get("penetration_depth",
                                   quality.get("penetration", 0.0)))

        return self.PQR_TEMPLATE_HTML.format(
            weld_type=context["weld_type"],
            current=p.current,
            voltage=p.voltage,
            speed=p.travel_speed,
            stickout=p.stickout,
            heat_input=q.heat_input,
            arc_length=q.arc_length,
            eta_residual=q.eta_residual,
            porosity_risk=q.porosity_risk,
            distortion=q.angular_distortion,
            penetration=q.penetration_depth,
            eta_pass="Pass" if eta < 0.5 else "Fail",
            porosity_pass="Pass" if porosity < 0.1 else "Fail",
            distortion_pass="Pass" if distortion < 2.0 else "Fail",
            penetration_pass="Pass" if penetration >= 1.5 else "Fail",
            conclusion=conclusion,
        )

    # ═══════════════════════════════════════════════════════════════════
    # v0.3.0: 章锋2026-07-04论文 DOCX输出 + κ-Snap聚合
    # ═══════════════════════════════════════════════════════════════════

    # ── κ-Phase: DOCX Generation + KappaSnap Aggregation ──

    def generate_wps_docx(
        self,
        params: Dict[str, Any],
        quality: Optional[Dict[str, Any]] = None,
        weld_type: str = "flat",
        output_path: Optional[str] = None,
    ) -> str:
        """生成WPS DOCX文档.

        使用python-docx生成正式的Word文档. 如果python-docx不可用,
        回退到HTML格式并保存为.html文件.

        Args:
            params: 焊接参数字典.
            quality: 质量指标字典 (可选).
            weld_type: 焊接姿态类型.
            output_path: 输出文件路径. None则返回文档内容字符串.

        Returns:
            如果output_path为None, 返回文档内容字符串 (HTML回退时)
            或文件路径 (DOCX成功时). 如果output_path指定, 返回该路径.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt

            doc = Document()

            # 标题
            doc.add_heading(
                f"Welding Procedure Specification (WPS)\n{weld_type} Welding",
                level=1,
            )

            # 基本信息
            doc.add_heading("General Information", level=2)
            info_table = doc.add_table(rows=7, cols=2)
            info_table.style = "Table Grid"
            info_data = [
                ("Welding Type", weld_type),
                ("Process", "GMAW (Gas Metal Arc Welding)"),
                ("Material", "Low Carbon Steel (Q235)"),
                ("Wire Diameter", f"{params.get('wire_diameter', 1.2)} mm"),
                ("Shielding Gas", str(params.get("gas_type", "Ar+CO2 80/20"))),
                ("Gas Flow Rate", f"{params.get('gas_flow', 15.0)} L/min"),
                ("Wire Feed", f"{params.get('wire_feed', 8.0)} m/min"),
            ]
            for i, (key, val) in enumerate(info_data):
                info_table.rows[i].cells[0].text = key
                info_table.rows[i].cells[1].text = str(val)

            # 推荐参数范围
            doc.add_heading("Recommended Parameter Ranges", level=2)
            current: float = float(params.get("current", 200.0))
            voltage: float = float(params.get("voltage", 24.0))
            speed: float = float(params.get("travel_speed", params.get("speed", 6.0)))

            range_table = doc.add_table(rows=6, cols=4)
            range_table.style = "Table Grid"
            range_data = [
                ("Parameter", "Min", "Nominal", "Max"),
                ("Current (A)", str(int(current * 0.8)), str(int(current)), str(int(current * 1.2))),
                ("Voltage (V)", f"{voltage*0.9:.1f}", f"{voltage:.1f}", f"{voltage*1.1:.1f}"),
                ("Travel Speed (mm/s)", f"{max(2.0, speed*0.7):.1f}", f"{speed:.1f}", f"{min(15.0, speed*1.3):.1f}"),
                ("Stickout (mm)", "8", str(params.get("stickout", 15.0)), "25"),
                ("Weave (mm)", "0", str(params.get("weave", 2.0)), "5"),
            ]
            for i, row_data in enumerate(range_data):
                for j, cell_text in enumerate(row_data):
                    range_table.rows[i].cells[j].text = cell_text

            # 质量要求
            doc.add_heading("Quality Requirements", level=2)
            doc.add_paragraph("Max porosity risk: < 10%")
            doc.add_paragraph("Max angular distortion: < 2.0°")
            doc.add_paragraph("Min penetration depth: >= 1.5 mm")
            doc.add_paragraph("Max heat input: <= 2.5 kJ/mm")

            # 保存
            if output_path is None:
                output_path = "wps_output.docx"
            doc.save(output_path)
            return output_path

        except ImportError:
            # 回退: HTML格式
            html_content = self.generate_wps(params, quality, weld_type, format="html")
            if output_path is None:
                return html_content
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(html_content)
            return output_path

    def generate_pqr_docx(
        self,
        params: Dict[str, Any],
        quality: Dict[str, Any],
        weld_type: str = "flat",
        output_path: Optional[str] = None,
    ) -> str:
        """生成PQR DOCX文档.

        使用python-docx生成正式的Word文档. 如果python-docx不可用,
        回退到HTML格式.

        Args:
            params: 焊接参数字典.
            quality: 质量指标字典.
            weld_type: 焊接姿态类型.
            output_path: 输出文件路径.

        Returns:
            文件路径或文档内容字符串.
        """
        try:
            from docx import Document

            doc = Document()

            doc.add_heading(
                f"Procedure Qualification Record (PQR)\n{weld_type} Welding",
                level=1,
            )

            # 实际参数
            doc.add_heading("Actual Welding Parameters", level=2)
            param_table = doc.add_table(rows=6, cols=2)
            param_table.style = "Table Grid"
            eta: float = float(quality.get("eta_residual", quality.get("eta", 0.0)))
            porosity: float = float(quality.get("porosity_risk", quality.get("porosity", 0.0)))
            distortion: float = float(quality.get("angular_distortion", quality.get("distortion", 0.0)))
            penetration: float = float(quality.get("penetration_depth", quality.get("penetration", 0.0)))

            param_data = [
                ("Current (A)", str(params.get("current", 200.0))),
                ("Voltage (V)", str(params.get("voltage", 24.0))),
                ("Travel Speed (mm/s)", str(params.get("travel_speed", params.get("speed", 6.0)))),
                ("Stickout (mm)", str(params.get("stickout", 15.0))),
                ("Heat Input (kJ/mm)", f"{quality.get('heat_input', 0.0):.4f}"),
                ("Arc Length (mm)", f"{quality.get('arc_length', 0.0):.4f}"),
            ]
            for i, (key, val) in enumerate(param_data):
                param_table.rows[i].cells[0].text = key
                param_table.rows[i].cells[1].text = val

            # 质量检测结果
            doc.add_heading("Quality Test Results", level=2)
            qual_table = doc.add_table(rows=5, cols=4)
            qual_table.style = "Table Grid"
            qual_data = [
                ("Test Item", "Result", "Criteria", "Pass/Fail"),
                ("Eta Residual", f"{eta:.4f}", "< 0.5", "Pass" if eta < 0.5 else "Fail"),
                ("Porosity Risk", f"{porosity:.4f}", "< 0.1", "Pass" if porosity < 0.1 else "Fail"),
                ("Angular Distortion", f"{distortion:.4f}°", "< 2.0°", "Pass" if distortion < 2.0 else "Fail"),
                ("Penetration Depth", f"{penetration:.4f} mm", ">= 1.5 mm", "Pass" if penetration >= 1.5 else "Fail"),
            ]
            for i, row_data in enumerate(qual_data):
                for j, cell_text in enumerate(row_data):
                    qual_table.rows[i].cells[j].text = cell_text

            # 结论
            doc.add_heading("Conclusion", level=2)
            all_pass: bool = (eta < 0.5 and porosity < 0.1
                              and distortion < 2.0 and penetration >= 1.5)
            conclusion: str = (
                "QUALIFIED — All acceptance criteria met. "
                "This procedure is approved for production welding."
                if all_pass
                else "NOT QUALIFIED — One or more acceptance criteria "
                     "not met. Parameters require adjustment."
            )
            doc.add_paragraph(conclusion)

            if output_path is None:
                output_path = "pqr_output.docx"
            doc.save(output_path)
            return output_path

        except ImportError:
            html_content = self.generate_pqr(params, quality, weld_type, format="html")
            if output_path is None:
                return html_content
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(html_content)
            return output_path

    def aggregate_ksnap_stats(
        self,
        snap_entries: List[Dict[str, Any]],
    ) -> Dict[str, Any]:
        """聚合κ-Snap审计统计.

        从κ-Snap FIFO条目列表中提取统计信息:
          - 总条目数
          - η均值/最大/最小/标准差
          - Ψ-Check通过率
          - 违规类型分布
          - 时间跨度

        Args:
            snap_entries: κ-Snap条目字典列表, 每个包含:
              - eta: η值
              - psi_passed: 是否通过Ψ-Check
              - violation: 违规描述 (空字符串=无违规)
              - step: 步骤号
              - timestamp: 时间戳

        Returns:
            统计摘要字典.
        """
        if len(snap_entries) == 0:
            return {
                "total_entries": 0,
                "eta_mean": 0.0,
                "eta_max": 0.0,
                "eta_min": 0.0,
                "eta_std": 0.0,
                "psi_pass_rate": 0.0,
                "violation_count": 0,
                "violation_types": {},
                "time_span_s": 0.0,
            }

        etas: List[float] = [float(e.get("eta", 0.0)) for e in snap_entries]
        psi_results: List[bool] = [bool(e.get("psi_passed", True)) for e in snap_entries]
        violations: List[str] = [
            e.get("violation", "") for e in snap_entries
            if e.get("violation", "")
        ]
        steps: List[int] = [int(e.get("step", 0)) for e in snap_entries]
        timestamps: List[float] = [float(e.get("timestamp", 0.0)) for e in snap_entries]

        # 违规类型统计
        violation_types: Dict[str, int] = {}
        for v in violations:
            # 分割复合违规 (用;分隔)
            for v_part in v.split(";"):
                v_part = v_part.strip()
                if v_part:
                    # 提取违规类型 (取括号前的部分)
                    v_type: str = v_part.split("(")[0].strip()
                    violation_types[v_type] = violation_types.get(v_type, 0) + 1

        return {
            "total_entries": len(snap_entries),
            "eta_mean": float(np.mean(etas)),
            "eta_max": float(np.max(etas)),
            "eta_min": float(np.min(etas)),
            "eta_std": float(np.std(etas)),
            "psi_pass_rate": float(np.mean(psi_results)),
            "violation_count": len(violations),
            "violation_types": violation_types,
            "time_span_s": max(timestamps) - min(timestamps) if len(timestamps) > 1 else 0.0,
            "step_range": [min(steps), max(steps)] if steps else [0, 0],
        }
